import streamlit as st
import io
try:
    import docx
except ImportError:
    docx = None

class Helpers:
    """Helper utilities for file downloads and format conversions."""
    
    @staticmethod
    def download_file(content: str, filename: str, mime_type: str = "text/plain"):
        """Render a Streamlit download button for text content."""
        st.download_button(
            label=f"⬇️ Download {filename.split('.')[-1].upper()}",
            data=content.encode('utf-8'),
            file_name=filename,
            mime=mime_type,
            key=f"dl_{filename}"
        )
    
    @staticmethod
    def download_docx(content: str, filename: str = "email.docx"):
        """Render a Streamlit download button for DOCX format."""
        if docx is None:
            st.error("python-docx package is not installed.")
            return
            
        doc = docx.Document()
        lines = content.split('\n')
        for line in lines:
            if line.startswith("Subject:"):
                h = doc.add_heading(line, level=1)
            else:
                doc.add_paragraph(line)
                
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.download_button(
            label="⬇️ Download DOCX",
            data=buffer,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"dl_docx_{filename}"
        )
